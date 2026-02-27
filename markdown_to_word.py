#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown文件转换为Word文档
"""

import os
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def convert_markdown_to_word(markdown_file, output_file):
    """
    将Markdown文件转换为Word文档
    
    Args:
        markdown_file: Markdown文件路径
        output_file: 输出Word文档路径
    """
    # 读取Markdown内容
    with open(markdown_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换为HTML
    html_content = markdown.markdown(md_content)
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档属性
    doc.core_properties.title = os.path.basename(output_file).replace('.docx', '')
    doc.core_properties.author = '系统开发团队'
    doc.core_properties.subject = '电子文件销售系统'
    doc.core_properties.category = '技术文档'
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    # 设置中文字体支持
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 转换HTML到Word
    # 这里我们使用简单的HTML解析，将内容分段添加到文档中
    paragraphs = html_content.split('\n')
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # 处理标题
        if para.startswith('<h1>') and para.endswith('</h1>'):
            title = para.replace('<h1>', '').replace('</h1>', '')
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.bold = True
        elif para.startswith('<h2>') and para.endswith('</h2>'):
            title = para.replace('<h2>', '').replace('</h2>', '')
            heading = doc.add_heading(title, level=2)
            heading.runs[0].font.size = Pt(14)
            heading.runs[0].font.bold = True
        elif para.startswith('<h3>') and para.endswith('</h3>'):
            title = para.replace('<h3>', '').replace('</h3>', '')
            heading = doc.add_heading(title, level=3)
            heading.runs[0].font.size = Pt(12)
            heading.runs[0].font.bold = True
        elif para.startswith('<h4>') and para.endswith('</h4>'):
            title = para.replace('<h4>', '').replace('</h4>', '')
            heading = doc.add_heading(title, level=4)
            heading.runs[0].font.size = Pt(11)
            heading.runs[0].font.bold = True
        
        # 处理表格
        elif para.startswith('<table>') and para.endswith('</table>'):
            # 简化表格处理，这里只添加一个说明
            doc.add_paragraph('(表格内容请参考原始Markdown文件)')
        
        # 处理代码块
        elif para.startswith('<pre>') and para.endswith('</pre>'):
            code = para.replace('<pre>', '').replace('</pre>', '')
            if code.startswith('<code>') and code.endswith('</code>'):
                code = code.replace('<code>', '').replace('</code>', '')
            p = doc.add_paragraph()
            p.add_run('代码块: ').bold = True
            p.add_run(code[:100] + '...' if len(code) > 100 else code)
            p.italic = True
        
        # 处理列表
        elif para.startswith('<ul>') or para.startswith('<ol>'):
            # 简化列表处理
            doc.add_paragraph('(列表内容请参考原始Markdown文件)')
        
        # 处理普通段落
        elif para.startswith('<p>') and para.endswith('</p>'):
            text = para.replace('<p>', '').replace('</p>', '')
            p = doc.add_paragraph(text)
        
        # 其他内容
        else:
            doc.add_paragraph(para)
    
    # 保存文档
    doc.save(output_file)
    print(f"已将 {markdown_file} 转换为 {output_file}")


def main():
    """
    主函数
    """
    # 定义输入输出文件
    files = [
        ('数据库说明书.md', '数据库说明书.docx'),
        ('详细设计说明书.md', '详细设计说明书.docx')
    ]
    
    # 当前目录
    current_dir = os.getcwd()
    
    # 转换所有文件
    for md_file, docx_file in files:
        md_path = os.path.join(current_dir, md_file)
        docx_path = os.path.join(current_dir, docx_file)
        
        if os.path.exists(md_path):
            convert_markdown_to_word(md_path, docx_path)
        else:
            print(f"文件 {md_path} 不存在")


if __name__ == '__main__':
    main()
